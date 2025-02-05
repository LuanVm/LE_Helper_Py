from PyQt6.QtCore import QThread, pyqtSignal
from docx import Document
from docx2pdf import convert
import os
import pythoncom
from validate_docbr import CPF, CNPJ

class ContratoWorker(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, template_path, data, output_dir):
        super().__init__()
        self.template_path = template_path
        self.data = data
        self.output_dir = output_dir

    def run(self):
        try:
            pythoncom.CoInitialize()
            # Etapa 1: Carregar template
            doc = Document(self.template_path)
            self.progress.emit(10)

            # Etapa 2: Substituir placeholders
            self.substituir_placeholders(doc)
            self.progress.emit(50)

            # Etapa 3: Salvar DOCX temporário
            temp_docx = os.path.join(self.output_dir, "temp_contrato.docx")
            doc.save(temp_docx)
            self.progress.emit(70)

            # Etapa 4: Converter para PDF
            pdf_path = os.path.join(self.output_dir, "Contrato_Assinado.pdf")
            convert(temp_docx, pdf_path)
            self.progress.emit(90)

            # Etapa 5: Limpeza
            os.remove(temp_docx)
            self.progress.emit(100)
            self.finished.emit(pdf_path)
        except Exception as e:
            self.error.emit(str(e))
        finally:
            pythoncom.CoUninitialize()

    def validar_cpf(self, cpf_value):
        """Remove formatação e valida o CPF."""
        validador = CPF()
        cpf_digits = ''.join(filter(str.isdigit, cpf_value))
        return validador.validate(cpf_digits)

    def validar_cnpj(self, cnpj_value):
        """Remove formatação e valida o CNPJ."""
        validador = CNPJ()
        cnpj_digits = ''.join(filter(str.isdigit, cnpj_value))
        return validador.validate(cnpj_digits)

    def substituir_placeholders(self, doc):
        """Substitui os placeholders do DOCX pelos valores fornecidos."""
        for paragraph in doc.paragraphs:
            for key, value in self.data.items():
                if key in paragraph.text:
                    if not value.strip():
                        value = "NÃO INFORMADO"
                    if key in ["rep1_cpf", "rep2_cpf"]:
                        if not self.validar_cpf(value):
                            value = "CPF INVÁLIDO"
                    if key == "cnpj":
                        if not self.validar_cnpj(value):
                            value = "CNPJ INVÁLIDO"
                    paragraph.text = paragraph.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in self.data.items():
                        if key in cell.text:
                            if not value.strip():
                                value = "NÃO INFORMADO"
                            if key in ["rep1_cpf", "rep2_cpf"]:
                                if not self.validar_cpf(value):
                                    value = "CPF INVÁLIDO"
                            if key == "cnpj":
                                if not self.validar_cnpj(value):
                                    value = "CNPJ INVÁLIDO"
                            cell.text = cell.text.replace(key, value)